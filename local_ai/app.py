from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import requests
import re
import pdfplumber
from docx import Document
import json

app = Flask(__name__)
CORS(app, origins=["http://localhost:3000", "http://localhost:5173", "http://localhost:5174"])

@app.route('/generate_summary', methods=['POST'])
def generate_summary():
    data = request.get_json()
    answers_and_scores = data.get('answersAndScores', [])
    # Compose Groq prompt for summary
    prompt = f"""
    Based on these 6 interview answers and scores, generate a 2-3 sentence summary of the candidate's overall performance, strengths, and areas for improvement. Format as plain text only.
    Answers and scores:
    {json.dumps(answers_and_scores, indent=2)}
    """
    try:
        response = requests.post(
            GROQ_API_URL,
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": MODEL_NAME,
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 256,
                "temperature": 0.5
            }
        )
        groq_data = response.json()
        content = groq_data["choices"][0]["message"]["content"]
        # Return summary as plain text
        return jsonify({"summary": content.strip()})
    except Exception as e:
        print("Error in /generate_summary:", e)
        return jsonify({'error': str(e)}), 500

@app.route('/validate_resume', methods=['POST'])
def validate_resume():
    if 'resume' not in request.files:
        return jsonify({'error': 'No resume file uploaded'}), 400
    resume_file = request.files['resume']
    filename = resume_file.filename
    text = ""
    # Extract text from PDF
    if filename.lower().endswith('.pdf'):
        with pdfplumber.open(resume_file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    # Extract text from DOCX
    elif filename.lower().endswith('.docx'):
        doc = Document(resume_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        return jsonify({'error': 'Unsupported file type'}), 400
    # Compose Groq prompt for intelligent validation
    prompt = f"""
    Analyze this document and respond with JSON:
    {{
      'is_valid_resume': true/false,
      'reason': 'explanation if invalid',
      'extracted_info': {{
        'name': 'found name or null',
        'email': 'found email or null',
        'phone': 'found phone or null'
      }},
      'missing_fields': ['list of missing required fields'],
      'confidence': 'high/medium/low'
    }}
    Document:
    {text}
    """
    try:
        response = requests.post(
            GROQ_API_URL,
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": MODEL_NAME,
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 512,
                "temperature": 0.3
            }
        )
        groq_data = response.json()
        content = groq_data["choices"][0]["message"]["content"]
        # Try to extract JSON from Groq response
        json_match = re.search(r'\{[\s\S]*\}', content)
        if json_match:
            try:
                json_string = json_match.group(0)
                # Remove JS-style comments and trailing commas
                json_string = re.sub(r'//.*?(?=\n|$)', '', json_string)
                json_string = re.sub(r',([\s]*[}}\]])', r'\1', json_string)
                result = json.loads(json_string.replace("'", '"'))
                return jsonify(result)
            except Exception as e:
                print("JSON decode error in /validate_resume:", e)
                return jsonify({'error': 'AI response could not be parsed.'}), 500
        return jsonify({'error': 'No valid AI response received.'}), 500
    except Exception as e:
        print("Error in /validate_resume:", e)
        return jsonify({'error': str(e)}), 500
@app.route('/test', methods=['GET'])
def test():
    return jsonify({'message': 'Flask is working'})

GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_API_KEY = (os.getenv("GROQ_API_KEY") or "").strip()
if not GROQ_API_KEY or GROQ_API_KEY == "YOUR_GROQ_API_KEY":
    raise RuntimeError("GROQ_API_KEY environment variable not set or contains invalid value. Please set it with your actual Groq API key.")
print(f"Groq API key loaded: {GROQ_API_KEY[:10]}... (set via env, stripped)")
MODEL_NAME = "llama-3.1-8b-instant"

@app.route('/extract_resume', methods=['POST'])
def extract_resume():
    import re
    import pdfplumber
    from docx import Document
    if 'resume' not in request.files:
        return jsonify({'error': 'No resume file uploaded'}), 400
    resume_file = request.files['resume']
    filename = resume_file.filename
    text = ""
    # Extract text from PDF
    if filename.lower().endswith('.pdf'):
        with pdfplumber.open(resume_file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    # Extract text from DOCX
    elif filename.lower().endswith('.docx'):
        doc = Document(resume_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        return jsonify({'error': 'Unsupported file type'}), 400
    # Robust resume validation
    required_keywords = [
        'name', 'email', 'phone', 'experience', 'skills', 'education', 'work', 'project', 'summary', 'profile', 'certification', 'degree', 'university', 'college', 'employment', 'responsibility', 'achievement', 'internship', 'company', 'role', 'position'
    ]
    found_keywords = [kw for kw in required_keywords if re.search(rf'\b{kw}\b', text, re.IGNORECASE)]
    # Structure pattern: look for sections like Experience, Education, Skills, Projects
    section_patterns = [r'Experience', r'Education', r'Skills?', r'Projects?', r'Summary', r'Profile']
    found_sections = [pat for pat in section_patterns if re.search(rf'\b{pat}\b', text, re.IGNORECASE)]
    # Basic length check
    if not text or len(text.strip()) < 100:
        return jsonify({'error': 'Resume document is too short or empty. Please upload a valid resume.'}), 400
    # Require at least 3 sections and 5 keywords
    if len(found_sections) < 3 or len(found_keywords) < 5:
        return jsonify({'error': 'Resume document does not appear to be a valid resume. Please upload a document with sections like Experience, Education, Skills, and typical resume keywords.'}), 400
    # Regex extraction
    email_match = re.search(r'[\w\.-]+@[\w\.-]+', text)
    phone_match = re.search(r'(\+?\d{1,2}[\s-]?)?(\(?\d{3}\)?[\s-]?)?\d{3}[\s-]?\d{4}', text)
    # Name extraction: naive, first non-empty line not containing email/phone
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    name = ""
    for l in lines:
        if (email_match and email_match.group() in l) or (phone_match and phone_match.group() in l):
            continue
        if len(l.split()) >= 2 and not re.search(r'\d', l):
            name = l
            break
    # Return full extracted text for downstream use
    return jsonify({
        'name': name or 'Unknown',
        'email': email_match.group() if email_match else 'Unknown',
        'phone': phone_match.group() if phone_match else 'Unknown',
        'resume': filename,
        'raw_text': text  # Pass full text for question generation
    })

@app.route('/generate_question', methods=['POST'])
def generate_question():
    data = request.get_json()
    # Accept either resume filename or full extracted text
    resume_text = data.get('resume', '')
    # If resume_text is a dict (from frontend), get raw_text
    if isinstance(resume_text, dict) and 'raw_text' in resume_text:
        resume_text = resume_text['raw_text']
    difficulty = data.get('difficulty', 'medium')
    # If resume_text looks like a filename, try to get the extracted text from the candidate info
    # Otherwise, use the raw text directly
    if os.path.exists(resume_text):
        with open(resume_text, 'r') as f:
            resume_text = f.read()
    pythonprompt = f"""
You are conducting a technical interview for a Full Stack Developer position (React/Node.js).

Generate ONE {difficulty} level question that can be answered verbally in the time limit:
- Easy (20 seconds): Simple concept explanation or definition
- Medium (60 seconds): How-to or comparison question  
- Hard (120 seconds): Design approach or problem-solving strategy

DIFFICULTY GUIDELINES:
Easy examples: "What is JSX in React?", "Explain what npm is", "What's the difference between let and const?"
Medium examples: "How would you handle form validation in React?", "Explain REST API best practices"
Hard examples: "How would you architect a scalable authentication system?", "Explain React performance optimization strategies"

IMPORTANT: 
- Keep questions SHORT and FOCUSED on ONE topic
- Questions should be ANSWERABLE in the time limit
- Do NOT ask for code implementation or multiple steps
- Do NOT ask for "code snippets" or "create/build" anything
- Focus on EXPLAINING concepts, approaches, or strategies

Return ONLY valid JSON (no comments):
{{
    "question": "single focused question here",
    "category": "React/Node.js/Full Stack",
    "difficulty": "{difficulty}",
    "expected_duration": {180 if difficulty == 'easy' else 300 if difficulty == 'medium' else 600}
}}
"""
    prompt = pythonprompt
    import re
    import json
    try:
        response = requests.post(
            GROQ_API_URL,
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": MODEL_NAME,
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 512,
                "temperature": 0.7
            }
        )
        groq_data = response.json()
        content = groq_data["choices"][0]["message"]["content"]
        print("Raw Groq response:", content)
        # Try to extract JSON only (ignore markdown, extra text)
        json_match = re.search(r'\{[\s\S]*\}', content)
        if json_match:
            try:
                json_string = json_match.group(0)
                # Remove JavaScript-style comments
                json_string = re.sub(r'//.*?(?=\n|$)', '', json_string)
                # Remove trailing commas before closing braces/brackets
                json_string = re.sub(r',([\s]*[}}\]])', r'\1', json_string)
                question_data = json.loads(json_string)
                return jsonify(question_data)
            except json.JSONDecodeError as e:
                print("JSONDecodeError:", e)
                # Fallback: return default question
                return jsonify({"question": "What is Python and how is it used in Machine Learning?", "difficulty": difficulty})
        # Fallback: try to parse any JSON in the content
        try:
            question_data = json.loads(content)
            return jsonify(question_data)
        except Exception as e:
            print("Fallback JSON parse error:", e)
        # If all else fails, return raw content
        return jsonify({"question": content.strip(), "difficulty": difficulty})
    except Exception as e:
        print("General error in /generate_question:", e)
        return jsonify({'error': str(e)}), 500

@app.route('/evaluate_answer', methods=['POST'])
def evaluate_answer():
    data = request.json if request.json else request.get_json()
    question = data.get('question')
    answer = data.get('answer')
    difficulty = data.get('difficulty')
    # If answer is blank, return score 0 and feedback immediately
    if not answer or str(answer).strip() == "":
        print("Blank answer detected. Returning score 0 and feedback without Groq call.")
        return jsonify({"score": 0, "feedback": "No answer provided"})
    print("=" * 50)
    print("EVALUATE_ANSWER ENDPOINT CALLED")
    data = request.json if request.json else request.get_json()
    print(f"Received data: {data}")
    question = data.get('question')
    answer = data.get('answer')
    difficulty = data.get('difficulty')
    print(f"Question: {question}")
    print(f"Answer: {answer}")
    print(f"Difficulty: {difficulty}")
    prompt = f"""
You are an expert technical interviewer. Score this answer strictly:

Question: {question}
Answer: {answer}

Scoring rules:
- Only give a score of 8-10 for truly excellent, complete, and correct answers.
- Give 4-7 for partial, incomplete, or somewhat correct answers.
- Give 0-3 for mostly incorrect, missing, or irrelevant answers.
- Be realistic and do not give high scores for partial or vague answers.

Return ONLY this JSON (keep feedback under 100 words):
{{"score": 0-10, "feedback": "brief feedback"}}
"""
    try:
        print("Calling Groq API...")
        response = requests.post(
            GROQ_API_URL,
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": MODEL_NAME,
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 256,
                "temperature": 0.1
            }
        )
        print(f"Groq raw response: {response}")
        groq_data = response.json()
        print(f"Groq message content: {groq_data}")
        finish_reason = groq_data["choices"][0].get("finish_reason")
        if finish_reason == 'length':
            print("Groq response was truncated. Returning default score for blank answer.")
            return jsonify({"score": 0, "feedback": "No answer provided"})
        # Try to extract score and feedback
        score = 0
        feedback = "Unable to evaluate"
        try:
            content = groq_data["choices"][0]["message"]["content"]
            import re, json
            json_match = re.search(r'\{[\s\S]*\}', content)
            if json_match:
                json_string = json_match.group(0)
                json_string = re.sub(r'//.*?(?=\n|$)', '', json_string)
                json_string = re.sub(r',([\s]*[}}\]])', r'\1', json_string)
                result = json.loads(json_string)
                score = result.get('score', 0)
                feedback = result.get('feedback', 'Unable to evaluate')
        except Exception as e:
            print(f"Error extracting score/feedback: {e}")
            score = 0
            feedback = "Evaluation error"
        print(f"EXTRACTED SCORE: {score}")
        print(f"EXTRACTED FEEDBACK: {feedback}")
        print(f"RETURNING TO FRONTEND: {{'score': {score}, 'feedback': {feedback}}}")
        print("=" * 50)
        return jsonify({"score": score, "feedback": feedback})
    except Exception as e:
        print(f"ERROR IN EVALUATION: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/get_candidate', methods=['GET'])
def get_candidate():
    # Dummy candidate data (replace with real DB lookup if needed)
    return jsonify({
        'name': 'John Doe',
        'email': 'john@example.com',
        'phone': '123-456-7890',
        'resume': 'resume.pdf',
        'interview_progress': 'incomplete'
    })

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    app.run(host='0.0.0.0', port=port, debug=True)
